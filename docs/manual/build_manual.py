# -*- coding: utf-8 -*-
"""Genera el Manual del Operador de Hospiwaste (.docx) con capturas embebidas.
Texto en español de Venezuela (tratamiento de "tú")."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as _PILImage

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "screenshots")

NAVY = RGBColor(0x0B, 0x1A, 0x48)
ACCENT = RGBColor(0x2A, 0x27, 0xE9)
GREY = RGBColor(0x68, 0x68, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Segoe UI"
normal.font.size = Pt(10.5)
normal.font.color.rgb = NAVY

for lvl, size in [("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 11.5)]:
    st = doc.styles[lvl]
    st.font.name = "Segoe UI Semibold"
    st.font.size = Pt(size)
    st.font.color.rgb = NAVY if lvl != "Heading 1" else ACCENT
    st.font.bold = True
doc.styles["Heading 1"].paragraph_format.page_break_before = True

for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def p(text="", size=10.5, bold=False, color=NAVY, align=None, italic=False, space_after=6):
    par = doc.add_paragraph()
    if align: par.alignment = align
    par.paragraph_format.space_after = Pt(space_after)
    if text:
        r = par.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        r.font.color.rgb = color
    return par


def bullet(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = par.add_run(bold_prefix); r.bold = True; r.font.color.rgb = NAVY
        r2 = par.add_run(text); r2.font.color.rgb = NAVY
    else:
        r = par.add_run(text); r.font.color.rgb = NAVY
    return par


def callout(title, items, fill, bar):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, fill)
    cell.width = Inches(6.6)
    tp = cell.paragraphs[0]
    tp.paragraph_format.space_after = Pt(3)
    tr = tp.add_run(title); tr.bold = True; tr.font.size = Pt(10.5)
    tr.font.color.rgb = RGBColor.from_string(bar)
    for it in items:
        ip = cell.add_paragraph()
        ip.paragraph_format.space_after = Pt(2)
        ip.paragraph_format.left_indent = Inches(0.12)
        run = ip.add_run("•  " + it); run.font.size = Pt(10)
        run.font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


SPLITDIR = os.path.join(HERE, "_assets", "_split")
os.makedirs(SPLITDIR, exist_ok=True)
MAX_H = 6.6


def _caption(text):
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    r = cp.add_run(text); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY


def img(name, width=2.5, caption=None):
    """Inserta una captura. Las muy altas se parten en dos mitades lado a lado
    para mantenerlas legibles sin desbordar la página."""
    path = os.path.join(SHOTS, name)
    w, h = _PILImage.open(path).size
    aspect = h / w
    if aspect <= 2.65:
        disp_w = min(width, MAX_H / aspect)
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_before = Pt(4); par.paragraph_format.space_after = Pt(2)
        par.add_run().add_picture(path, width=Inches(disp_w))
    else:
        im = _PILImage.open(path)
        cut_top = int(h * 0.53)
        cut_bot = int(h * 0.47)
        base = os.path.splitext(name)[0]
        pa = os.path.join(SPLITDIR, base + "_a.png")
        pb = os.path.join(SPLITDIR, base + "_b.png")
        im.crop((0, 0, w, cut_top)).save(pa)
        im.crop((0, cut_bot, w, h)).save(pb)
        half_aspect = cut_top / w
        col_w = min(2.55, MAX_H / half_aspect)
        tbl = doc.add_table(rows=1, cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, ph in ((tbl.cell(0, 0), pa), (tbl.cell(0, 1), pb)):
            cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.add_run().add_picture(ph, width=Inches(col_w))
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    if caption:
        _caption(caption + ("  (la captura se muestra en dos mitades: izquierda = parte superior, derecha = parte inferior)" if aspect > 2.65 else ""))


def imgs_row(items, width=2.3):
    n = len(items)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, cap) in enumerate(items):
        c = tbl.cell(0, i)
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _w, _h = _PILImage.open(os.path.join(SHOTS, name)).size
        disp = min(width, MAX_H / (_h / _w))
        pp.add_run().add_picture(os.path.join(SHOTS, name), width=Inches(disp))
        cc = tbl.cell(1, i).paragraphs[0]; cc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cc.add_run(cap); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def heading(text, level=1):
    return doc.add_heading(text, level=level)

# ════════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════════
cover = doc.add_table(rows=1, cols=1); cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cc = cover.cell(0, 0); shade(cc, "0B1A48"); cc.width = Inches(6.6)
cc.paragraphs[0].paragraph_format.space_before = Pt(28)
r = cc.paragraphs[0].add_run("HOSPIWASTE"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = WHITE
cc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = cc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Trazabilidad de Desechos Clínicos"); sr.font.size = Pt(12); sr.font.color.rgb = RGBColor(0xC9, 0xD3, 0xF0)
sub.paragraph_format.space_after = Pt(28)

doc.add_paragraph().paragraph_format.space_after = Pt(14)
p("Manual del Operador", size=24, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
p("Guía de uso de las pantallas: Inicio, Recorrido, Pesaje y Tratamiento",
  size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
p("Versión 1.0 — Junio 2026", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Dirigido al personal operativo de planta y recorrido", size=10, color=GREY,
  align=WD_ALIGN_PARAGRAPH.CENTER)

# ════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
_ct = doc.add_paragraph(); _ct.paragraph_format.space_after = Pt(8)
_ctr = _ct.add_run("Contenido"); _ctr.bold = True; _ctr.font.size = Pt(17); _ctr.font.color.rgb = ACCENT
par = doc.add_paragraph()
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
run_in = OxmlElement("w:r"); t_in = OxmlElement("w:t")
t_in.text = "Abre este documento en Word y haz clic derecho → Actualizar campos para ver el índice."
run_in.append(t_in); fld.append(run_in); par._p.append(fld)

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN Y ACCESO
# ════════════════════════════════════════════════════════════════════════════
heading("1. Introducción y acceso", level=1)
p("Hospiwaste es el sistema con el que registramos, paso a paso, el recorrido completo de "
  "los tachos de desechos clínicos: desde que se recogen sucios en las áreas del cliente, "
  "se pesan en planta y entran a cámara fría, hasta que se envían a tratamiento. Todo lo que "
  "registres queda guardado con fotos y horarios, y alimenta los reportes oficiales.")
p("Como operador vas a usar cuatro pantallas:", bold=True, space_after=4)
bullet("muestra el estado del día y los totales.", bold_prefix="Inicio (Dashboard): ")
bullet("registras la recogida de tachos sucios y la entrega de tachos limpios.", bold_prefix="Recorrido: ")
bullet("pesas cada tacho que llegó a planta y le tomas las fotos.", bold_prefix="Pesaje: ")
bullet("envías los tachos infecciosos de cámara fría al tratamiento final.", bold_prefix="Tratamiento: ")

p()
callout("Antes de empezar",
        ["Usa siempre tu propio usuario: cada registro queda asociado a quien lo hizo.",
         "Trabaja con buena señal. Si el celular pierde conexión, espera a reconectar antes de guardar.",
         "Ten el celular con batería: el recorrido y el pesaje usan cronómetro y cámara."],
        fill="E4F0F8", bar="0B1A48")

heading("1.1 Ingresar al sistema", level=2)
p("Abre la aplicación y escribe tu correo y contraseña. Toca el ícono del ojo si quieres "
  "verificar la contraseña antes de entrar. Luego presiona Ingresar.")
img("00-login.png", width=2.45, caption="Pantalla de ingreso. Escribe tu correo y contraseña y toca «Ingresar».")
callout("Errores que NO debes cometer",
        ["No compartas tu usuario ni dejes la sesión abierta en un celular ajeno.",
         "Si ves «Correo o contraseña incorrectos», revisa mayúsculas y espacios; no insistas a ciegas."],
        fill="FBE9E9", bar="B91C1C")

heading("1.2 Cómo moverte (barra inferior)", level=2)
p("En la parte de abajo siempre tienes la barra de navegación con Inicio, Recorrido, Pesaje y "
  "Reportes. El botón «Más» abre el resto de opciones, entre ellas Tratamiento.")
callout("Dónde está cada cosa",
        ["Inicio, Recorrido y Pesaje están directos en la barra de abajo.",
         "Tratamiento está dentro del botón «Más» (••• a la derecha).",
         "Para salir, abre «Más» y toca «Cerrar sesión»."],
        fill="E4F0F8", bar="0B1A48")

# ════════════════════════════════════════════════════════════════════════════
# 2. DASHBOARD
# ════════════════════════════════════════════════════════════════════════════
heading("2. Inicio (Dashboard)", level=1)
p("Es la primera pantalla al entrar. No se registra nada aquí: sirve para ver de un vistazo "
  "cómo va el día y detectar trabajo pendiente.")
img("01-dashboard.png", width=2.45, caption="Pantalla de Inicio: tarjetas de resumen y gráficos del día y del mes.")

heading("2.1 Qué dice cada tarjeta", level=2)
bullet("cuántos recorridos se finalizaron hoy.", bold_prefix="Recorridos hoy: ")
bullet("total de tachos que están en uso (en planta, en cliente, etc.).", bold_prefix="Tachos en circulación: ")
bullet("tachos que ya se recogieron en un recorrido pero todavía no se pesaron. Esta es tu cola de trabajo en Pesaje.", bold_prefix="Pendientes de pesar: ")
bullet("tachos infecciosos enviados a tratamiento.", bold_prefix="En tratamiento: ")

heading("2.2 Los gráficos", level=2)
bullet("dónde están los tachos en este momento (en planta vs. en cliente).", bold_prefix="Tachos en circulación: ")
bullet("compara lo procesado contra lo que falta procesar hoy.", bold_prefix="Kilogramos del día: ")
bullet("kilos recibidos y procesados por empresa. Puedes cambiar el mes con las flechas.", bold_prefix="Kilogramos del mes: ")
callout("Buenas prácticas",
        ["Mira «Pendientes de pesar» al llegar: ese número es lo que tienes que pesar.",
         "Si «Recorridos hoy» no refleja un recorrido que cerraste, revisa tu conexión.",
         "Usa el selector de mes para comparar, pero recuerda que no cambia ningún dato: solo muestra."],
        fill="E8F5EC", bar="15803D")

# ════════════════════════════════════════════════════════════════════════════
# 3. RECORRIDO
# ════════════════════════════════════════════════════════════════════════════
heading("3. Recorrido", level=1)
p("En el recorrido registras lo que pasa en terreno: qué tachos sucios recoges (que después "
  "irán a pesaje) y qué tachos limpios entregas al cliente. Cada recorrido lleva su empresa, "
  "su ubicación y sus fotos.")

heading("3.1 Elegir el tipo de recorrido", level=2)
p("Al entrar a Recorrido eliges entre recorrido de andén (el habitual, con 6 horarios fijos al "
  "día para desechos peligrosos) o recorrido de Morgue (sin horario fijo, según lo pida la operación).")
imgs_row([("10-recorrido-tipos.png", "Elige andén o Morgue."),
          ("11-recorrido-horarios.png", "Los 6 horarios del día.")], width=2.4)
p("En andén verás los 6 horarios. Toca el horario que estás haciendo. Un horario ya completado "
  "queda marcado y no se puede repetir hasta el día siguiente.")

heading("3.2 Seleccionar empresa e iniciar", level=2)
p("Antes de registrar nada, selecciona la Empresa del recorrido y toca Iniciar recorrido. Solo "
  "entonces se activa el cronómetro y se habilita el formulario.")
img("12-recorrido-empresa.png", width=2.45, caption="Paso 1: selecciona la empresa y toca «Iniciar recorrido».")
img("13-recorrido-en-curso.png", width=2.45, caption="Paso 2: recorrido en curso, formulario habilitado.")

heading("3.3 Agregar tachos sucios y limpios", level=2)
p("Toca «Agregar tachos sucios» para abrir la lista. Busca por número y toca cada tacho que "
  "recoges; quedan marcados. Cuando termines, toca «Listo». Haz lo mismo en «Agregar tachos "
  "limpios» para los que dejas lavados al cliente.")
img("14-recorrido-picker.png", width=2.45, caption="Selector de tachos: busca por número y toca los que correspondan.")
callout("Qué seleccionar",
        ["Sucios recogidos: todos los tachos llenos que te llevas a planta. Estos aparecerán luego en Pesaje.",
         "Limpios entregados: solo tachos vacíos y lavados que dejas al cliente.",
         "Un mismo tacho no puede estar en las dos listas a la vez."],
        fill="E4F0F8", bar="0B1A48")

heading("3.4 Ubicación, fotos y guardar el andén", level=2)
p("Completa Piso, Área y Andén para ubicar el recorrido. Toma al menos una foto (es obligatoria) "
  "tocando «Tomar foto». Luego toca «Guardar andén y agregar otro». Puedes registrar varios "
  "andenes dentro del mismo horario repitiendo el formulario.")
img("15-recorrido-form-lleno.png", width=2.45, caption="Andén con tachos, ubicación y foto cargados, listo para guardar.")

heading("3.5 Finalizar el recorrido", level=2)
p("Cuando hayas registrado todos los andenes del horario, toca Finalizar recorrido y confirma. "
  "Una vez finalizado, ese horario del día no se puede volver a abrir.")
callout("Buenas prácticas",
        ["Toma la foto donde se vean los tachos y el área; es el respaldo del recorrido.",
         "Registra el andén apenas lo termines: las fotos se suben en ese momento.",
         "Verifica los números de tacho antes de guardar; un número mal cargado descuadra el pesaje."],
        fill="E8F5EC", bar="15803D")
callout("Errores que NO debes cometer",
        ["No finalices el recorrido si todavía te falta registrar un andén: no podrás reabrirlo hoy.",
         "No metas un tacho sucio en la lista de limpios (ni al revés): cambia toda la trazabilidad.",
         "No olvides seleccionar la empresa: sin empresa no se puede iniciar."],
        fill="FBE9E9", bar="B91C1C")

# ════════════════════════════════════════════════════════════════════════════
# 4. PESAJE
# ════════════════════════════════════════════════════════════════════════════
heading("4. Pesaje", level=1)
p("En Pesaje registras cada tacho sucio que llegó a planta: su peso, su tipo de desecho y las "
  "fotos de la balanza y del tacho. El sistema calcula solo el peso neto restando la tara.")

heading("4.1 Iniciar la sesión de pesaje", level=2)
p("El formulario arranca bloqueado. Toca Iniciar pesaje para empezar el cronómetro y habilitarlo. "
  "Una sola sesión te permite pesar varios tachos seguidos.")
img("20-pesaje-inicio.png", width=2.45, caption="Formulario bloqueado: toca «Iniciar pesaje».")
img("21-pesaje-en-curso.png", width=2.45, caption="Sesión en curso con la lista de pendientes por pesar.")
p("Arriba vas a ver los tachos pendientes por pesar (los que recogiste en el recorrido). Si un "
  "tacho de esa lista hoy no llegó, toca «ausente» para sacarlo de la cola.")

heading("4.2 Completar el pesaje de un tacho", level=2)
p("Sigue este orden:")
bullet("selecciona el tacho de la lista de pendientes.", bold_prefix="1) Número de tacho: ")
bullet("normalmente «Peligroso infeccioso». Cámbialo solo si el desecho es otro (citotóxico, líquidos, metálicos, etc.).", bold_prefix="2) Tipo de desecho: ")
bullet("escribe lo que marca la balanza. El Peso neto se calcula solo (peso menos la tara del tacho).", bold_prefix="3) Peso bruto: ")
bullet("primero la foto de la balanza, después la foto del tacho. Las dos son obligatorias.", bold_prefix="4) Fotos: ")
bullet("toca «Guardar y agregar otro» para pasar al siguiente tacho.", bold_prefix="5) Guardar: ")
img("22-pesaje-form-lleno.png", width=2.45, caption="Pesaje completo: tacho, tipo, peso neto calculado y las dos fotos con fecha y hora.")
callout("Opciones especiales",
        ["¿Es un pesaje de Yaris?: actívalo solo si la carga viene de un tacho Yaris.",
         "Tratar inmediatamente: márcalo si ese tacho infeccioso se trata de inmediato y no pasa por cámara fría.",
         "Observaciones: úsalo para notas útiles (ej.: «tacho con tapa dañada»)."],
        fill="E4F0F8", bar="0B1A48")

heading("4.3 Editar y finalizar", level=2)
p("Mientras la sesión está abierta puedes tocar un tacho ya pesado para corregirlo. Cuando hayas "
  "pesado todos los pendientes, toca Finalizar pesaje y confirma: los tachos pasan "
  "automáticamente a cámara fría y ya no se pueden editar.")
callout("Buenas prácticas",
        ["Verifica que el peso bruto sea mayor que la tara; si no, el sistema marca error.",
         "La foto de la balanza debe mostrar el número legible; la del tacho, su número.",
         "Pesa todos los pendientes antes de finalizar: el botón se habilita cuando la cola está en cero."],
        fill="E8F5EC", bar="15803D")
callout("Errores que NO debes cometer",
        ["No uses «Cancelar» por error: descarta toda la sesión y sus fotos.",
         "No inventes un peso si la balanza falla; resuelve la balanza primero.",
         "No marques «Tratar inmediatamente» salvo que realmente se trate de una vez."],
        fill="FBE9E9", bar="B91C1C")

# ════════════════════════════════════════════════════════════════════════════
# 5. TRATAMIENTO
# ════════════════════════════════════════════════════════════════════════════
heading("5. Tratamiento", level=1)
p("En Tratamiento envías al proceso final los tachos infecciosos que están en cámara fría. "
  "La pantalla muestra solo esos tachos: no tienes que buscarlos.")
img("30-tratamiento.png", width=2.45, caption="Pantalla de Tratamiento. Cuando no hay tachos infecciosos en cámara fría, muestra este mensaje.")

heading("5.1 Cómo enviar tachos a tratamiento", level=2)
p("Cuando hay tachos disponibles, cada uno aparece como una tarjeta con su número y su tamaño. "
  "El proceso es:")
bullet("toca cada tacho que vas a tratar; queda marcado con un check azul.", bold_prefix="1) Seleccionar: ")
bullet("puedes marcar varios a la vez.", bold_prefix="2) Repetir: ")
bullet("toca «Enviar N a tratamiento». El botón muestra cuántos seleccionaste.", bold_prefix="3) Enviar: ")
bullet("aparece la confirmación en verde de que se registraron.", bold_prefix="4) Listo: ")
p("Si la lista está vacía («No hay tachos infecciosos en cámara fría»), significa que todavía no "
  "hay nada para tratar: primero deben pesarse tachos infecciosos que entren a cámara fría.",
  italic=True, color=GREY)
callout("Buenas prácticas",
        ["Confirma el número de cada tacho antes de enviarlo.",
         "Envía en grupo los tachos que van juntos al mismo proceso.",
         "Si esperabas ver un tacho y no aparece, revisa que se haya pesado como «infeccioso»."],
        fill="E8F5EC", bar="15803D")
callout("Errores que NO debes cometer",
        ["No envíes un tacho que no se trató realmente: la acción queda registrada.",
         "No selecciones tachos «por si acaso»; envía solo los que van a tratamiento ahora."],
        fill="FBE9E9", bar="B91C1C")

# ════════════════════════════════════════════════════════════════════════════
# 6. RESUMEN DEL FLUJO + GLOSARIO
# ════════════════════════════════════════════════════════════════════════════
heading("6. Resumen del flujo diario", level=1)
for i, txt in enumerate([
    "Ingresas con tu usuario y revisas «Pendientes de pesar» en Inicio.",
    "Haces el Recorrido del horario: seleccionas empresa, recoges tachos sucios, entregas limpios, foto y finalizas.",
    "En Pesaje pesas cada tacho recogido, con tipo de desecho y fotos, y finalizas (pasan a cámara fría).",
    "En Tratamiento envías los tachos infecciosos de cámara fría al proceso final.",
], start=1):
    bullet(txt, bold_prefix=f"Paso {i}: ")

p()
heading("7. Glosario rápido", level=1)
for term, desc in [
    ("Tacho", "el contenedor de desechos. Se identifica por número (ej.: 001)."),
    ("Tara", "el peso del tacho vacío. El sistema lo resta para dar el peso neto."),
    ("Peso neto", "el peso real del desecho = peso bruto − tara. Lo calcula el sistema."),
    ("Andén", "cada punto/registro dentro de un horario de recorrido."),
    ("Cámara fría", "donde quedan los tachos pesados a la espera de tratamiento."),
    ("Empresa", "el cliente/empresa al que pertenece el recorrido y sus tachos."),
]:
    bullet(desc, bold_prefix=f"{term}: ")

p()
p("Ante cualquier duda o si algo no se guarda, avisa a tu supervisor antes de seguir. Es "
  "preferible registrar bien y un poco más lento que arrastrar un error en la trazabilidad.",
  italic=True, color=GREY)

# ── Pie de página ──────────────────────────────────────────────────────────────
footer = doc.sections[0].footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Hospiwaste — Manual del Operador · v1.0 · Junio 2026")
fr.font.size = Pt(8); fr.font.color.rgb = GREY

out = os.path.join(HERE, "Manual_Operador_Hospiwaste.docx")
doc.save(out)
print("OK ->", out)
